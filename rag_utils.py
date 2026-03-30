from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import os
from pathlib import Path
from typing import Any, Iterable

from langchain_core.documents import Document

try:
    from langchain_community.vectorstores import FAISS
except ModuleNotFoundError:  # pragma: no cover - depends on optional install
    FAISS = None

try:
    from langchain_google_genai import GoogleGenerativeAIEmbeddings
except ModuleNotFoundError:  # pragma: no cover - depends on optional install
    GoogleGenerativeAIEmbeddings = None

try:
    from pypdf import PdfReader
except ModuleNotFoundError:  # pragma: no cover - depends on optional install
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ModuleNotFoundError:  # pragma: no cover - depends on optional install
    DocxDocument = None

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ModuleNotFoundError:  # pragma: no cover - depends on optional install
    RecursiveCharacterTextSplitter = None


@dataclass
class LoadedSource:
    source: str
    text: str


DEFAULT_EMBEDDING_MODEL_OPTIONS = [
    "gemini-embedding-001",
    "gemini-embedding-002",
]


def _read_txt_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore").strip()


def _read_pdf_bytes(data: bytes) -> str:
    if PdfReader is None:
        raise ModuleNotFoundError("pypdf is required to ingest PDF documents.")

    reader = PdfReader(BytesIO(data))
    text_parts = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts).strip()


def _read_docx_bytes(data: bytes) -> str:
    if DocxDocument is None:
        raise ModuleNotFoundError("python-docx is required to ingest DOCX documents.")

    doc = DocxDocument(BytesIO(data))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def read_source_bytes(name: str, data: bytes) -> LoadedSource | None:
    suffix = Path(name).suffix.lower()
    if suffix == ".txt":
        text = _read_txt_bytes(data)
    elif suffix == ".pdf":
        text = _read_pdf_bytes(data)
    elif suffix == ".docx":
        text = _read_docx_bytes(data)
    else:
        return None

    if not text:
        return None
    return LoadedSource(source=name, text=text)


def load_local_sources(knowledge_dir: Path) -> list[LoadedSource]:
    sources: list[LoadedSource] = []
    for path in sorted(knowledge_dir.glob("*")):
        if not path.is_file():
            continue
        loaded = read_source_bytes(path.name, path.read_bytes())
        if loaded:
            sources.append(loaded)
    return sources


def build_documents(sources: Iterable[LoadedSource]) -> list[Document]:
    docs: list[Document] = []
    if RecursiveCharacterTextSplitter is not None:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=150,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        split = splitter.split_text
    else:
        def split(text: str) -> list[str]:
            chunks: list[str] = []
            start = 0
            while start < len(text):
                end = min(len(text), start + 800)
                chunks.append(text[start:end].strip())
                if end >= len(text):
                    break
                start = max(end - 150, start + 1)
            return [chunk for chunk in chunks if chunk]

    for source in sources:
        for index, chunk in enumerate(split(source.text)):
            docs.append(
                Document(
                    page_content=chunk,
                    metadata={"source": source.source, "chunk": index},
                )
            )
    return docs


def get_available_embedding_models() -> list[str]:
    env_value = os.getenv("GEMINI_AVAILABLE_EMBEDDING_MODELS", "")
    configured = [item.strip() for item in env_value.split(",") if item.strip()]
    models = configured or DEFAULT_EMBEDDING_MODEL_OPTIONS
    return list(dict.fromkeys(models))


def parse_embedding_chain(
    primary_model: str | None = None, fallback_models: list[str] | None = None
) -> list[str]:
    primary = primary_model or os.getenv("GEMINI_EMBEDDING_MODEL", DEFAULT_EMBEDDING_MODEL_OPTIONS[0])
    configured = fallback_models or []
    if not configured:
        env_value = os.getenv("GEMINI_FALLBACK_EMBEDDING_MODELS", "")
        configured = [item.strip() for item in env_value.split(",") if item.strip()]

    chain: list[str] = []
    for candidate in [primary, *configured, *get_available_embedding_models()]:
        if candidate and candidate not in chain:
            chain.append(candidate)
    return chain


def is_embedding_limit_error(exc: Exception) -> bool:
    message = str(exc).lower()
    signals = [
        "429",
        "quota",
        "rate limit",
        "resource exhausted",
        "resource_exhausted",
        "too many requests",
        "exceeded",
        "limit reached",
    ]
    return any(signal in message for signal in signals)


def build_embeddings(model_name: str, api_key: str):
    if FAISS is None or GoogleGenerativeAIEmbeddings is None:
        raise ModuleNotFoundError(
            "langchain-community and langchain-google-genai are required to build the vector store."
        )

    return GoogleGenerativeAIEmbeddings(
        model=model_name,
        google_api_key=api_key,
    )


def build_vector_store(
    documents: list[Document],
    api_key: str,
    embedding_chain: list[str] | None = None,
    embedding_builder=build_embeddings,
) -> tuple[Any, str, list[str]]:
    if FAISS is None:
        raise ModuleNotFoundError(
            "langchain-community is required to build the vector store."
        )

    chain = embedding_chain or parse_embedding_chain()
    if not chain:
        raise ValueError("Embedding model chain is empty.")

    errors: list[str] = []
    last_exc: Exception | None = None
    for index, model_name in enumerate(chain):
        try:
            embeddings = embedding_builder(model_name, api_key)
            store = FAISS.from_documents(documents, embedding=embeddings)
            return store, model_name, errors
        except Exception as exc:  # noqa: BLE001
            last_exc = exc
            if index == len(chain) - 1 or not is_embedding_limit_error(exc):
                raise
            errors.append(f"{model_name}: {exc}")

    if last_exc is not None:
        raise last_exc
    raise RuntimeError("Vector store build failed without an exception.")


def retrieve_context(store: Any, query: str, top_k: int = 4) -> list[Document]:
    return store.similarity_search(query, k=top_k)
